try:
    from scripts.objects import mockStudent, json_to_student
    from scripts.fileSystemInteraction import getDirectory
    from scripts.settings import get_setting
    from scripts.helpers import get_naming_convention
except Exception as e:
    from objects import mockStudent, json_to_student
    from fileSystemInteraction import getDirectory
    from settings import get_setting  
    from helpers import get_naming_convention
from re import findall
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION

def doAuditMethod(studentObject):

    if (studentObject == 'mock'):
        studentObject = mockStudent()
    else:
        studentObject = json_to_student(studentObject)

    default_path = get_setting("default-path-for-audit")
    file_path = getDirectory('Save Audit Report', default_path)
    if not file_path:
        return

    destination = file_path + '/' + get_naming_convention(studentObject.name) + '_Audit.docx'
    try:
        generateAudit(studentObject, destination)
    except Exception as e:
        print(e)
        raise Exception("Error: Error at audit generation. Please try again later.")
    return destination

def generateAudit(studentObject, destination):
    # Class tracking for GPA
    core_complete = [course for course in studentObject.classes if course.grade and course.grade[0] != 'W' and (course.type == 'core' or course.type == 'following')]
    core_incomplete = [course for course in studentObject.classes if (not course.grade or course.grade[0] == 'W') and course.semester and (course.type == 'core' or course.type == 'following')]

    elective_complete = [course for course in studentObject.classes if course.grade and course.grade[0] != 'W' and (course.type == 'electives' or course.type == 'additional')]
    elective_incomplete = [course for course in studentObject.classes if (not course.grade or course.grade[0] == 'W') and (course.type == 'electives' or course.type == 'additional')]

    total_complete = [course for course in studentObject.classes if course.grade and course.grade[0] != 'W' and (is5000Level(course))]

    total_incomplete = [course for course in studentObject.classes if (not course.grade or course.grade[0] == 'W') and (course.type == 'core' or course.type == 'additional' or course.type == 'electives' or (course.type == 'following' and course.semester))]

    leveling_courses = [course for course in studentObject.classes if course.leveling]

    # Sort all courses by course number
    total_complete.sort(key=lambda x: x.number)
    total_incomplete.sort(key=lambda x: x.number)


    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # title
    para = doc.add_paragraph()
    para.add_run('Audit Report\n').bold = True
    para.alignment = 1

    # basic information
    para = doc.add_paragraph()

    inch = 914400
    tab_stops = para.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(round(inch*3.5))

    para.add_run('Name: ').bold = True
    para.add_run(studentObject.name)
    para.add_run('\tID: ').bold = True
    para.add_run(str(studentObject.studentId))
    para.add_run('\nPlan: ').bold = True
    para.add_run('Master')
    para.add_run('\tMajor: ').bold = True
    if studentObject.track == 'Software Engineering':
        para.add_run('Software Engineering')
    else:
        para.add_run('Computer Science')
    para.add_run('\n\tTrack: ').bold = True
    para.add_run(studentObject.track)


    #GPA
    core_complete.sort(key=lambda x: x.grade)
    elective_complete.sort(key=lambda x: x.grade)
    total_complete.sort(key=lambda x: x.grade)

    core_complete = core_complete[:5]
    elective_complete = elective_complete[:7]

    core_gpa = getGPA(core_complete)
    elective_gpa = getGPA(elective_complete)
    combined_gpa = getGPA(total_complete)

    para = doc.add_paragraph()
    para.add_run('\nCore GPA: ').bold = True
    para.add_run('%.3f' % core_gpa)
    para.add_run('\nElective GPA: ').bold = True
    para.add_run('%.3f' % elective_gpa)
    para.add_run('\nCombined GPA: ').bold = True
    para.add_run('%.3f' % combined_gpa)

    # courses
    para = doc.add_paragraph()
    para.add_run('\nCore Courses: ').bold = True

    cores = [course for course in studentObject.classes if course.type == 'core']
    cores = cores + [course for course in studentObject.classes if course.semester and course.type == 'following']
    para.add_run(", ".join([course.number.split("(")[0].strip() for course in cores]))


    para.add_run('\nElective Courses: ').bold = True

    electives = elective_complete + elective_incomplete
    electives.sort(key=lambda x: x.number)
    para.add_run(", ".join([course.number.split("(")[0].strip() for course in electives]))

    # leveling courses (incomplete) NEED LIST OF LEVELING COURSES HERE
    para = doc.add_paragraph()
    para.add_run('\nLeveling Courses and Pre-requisites from Admission Letter:\n').bold = True
    
    for course in leveling_courses:
        c = "".join(findall("[a-zA-Z]* *[0-9]*", course.number))
        if course.leveling == 'Completed':
            para.add_run("\n" + c + " - " + course.leveling + ": " + course.semester + ": " + course.grade)
        else:
            para.add_run("\n" + c + " - " + course.leveling)
    
    if (not leveling_courses):
        para.add_run("\nNone")

    # requirements (incomplete)
    core_status = ""
    elective_status = ""
    overall_status = ""

    para = doc.add_paragraph()
    para.add_run('\nOutstanding Requirements:\n').bold = True

    
    # Taking extra course?
    extra_course = False
    if (len(electives) >= 7):
        extra_course = True

    if((len(core_incomplete) == 0 and core_gpa >= 3.19) or (len(core_incomplete) == 0 and core_gpa >= 3.0 and extra_course)):
        core_status = "Core complete."
    elif(len(core_incomplete) == 0):
        core_status = "Core currently failing."
    elif(extra_course):
        core_status = printRequiredGPA(3.0, core_gpa, core_complete, core_incomplete)
        para.add_run('\nTo maintain a 3.0 core GPA:')
    else:
        core_status = printRequiredGPA(3.19, core_gpa, core_complete, core_incomplete)
        para.add_run('\nTo maintain a 3.19 core GPA:')
        
    para.add_run('\n' + core_status + ", ".join(course.number.split("(")[0].strip() for course in core_incomplete))
    
    if(len(elective_incomplete) == 0 and elective_gpa >= 3.0):
        elective_status = "Elective complete."
    elif(len(elective_incomplete) == 0 and elective_gpa < 3.0):
        elective_status = "Elective currently failing."
    else:
        elective_status = printRequiredGPA(3.0, elective_gpa, elective_complete, elective_incomplete)
        para.add_run('\nTo maintain a 3.0 elective GPA:')
    
    para.add_run('\n' + elective_status + ", ".join(course.number.split("(")[0].strip() for course in elective_incomplete))

    if(len(total_incomplete) == 0 and combined_gpa >= 3.0):
        overall_status = "Overall complete."
    elif(len(total_incomplete) == 0 and combined_gpa < 3.0):
        overall_status = "Overall currently failing."
    else:
        overall_status = printRequiredGPA(3.0, combined_gpa, total_complete, total_incomplete)
        para.add_run('\nTo maintain a 3.0 overall GPA:')
    
    para.add_run('\n' + overall_status + ", ".join(course.number.split("(")[0].strip() for course in total_incomplete))

    try:
        doc.save(destination)
    except PermissionError:
        raise Exception("Error: You do not have permissions to save here. Perhaps you have an older version open. Close out other applications and try again.")

def printRequiredGPA(required_GPA, GPA, completed_courses, remaining_courses):
    target = calculateRequiredGPA(required_GPA, GPA, completed_courses, remaining_courses)

    if target <= 2:
        return "\tThe student must pass: "
    
    if len(remaining_courses) == 1:
        grade = 'Error'

        if target > 4.0:
            return 'Impossible'
        elif target > 3.670:
            grade = 'A'
        elif target > 3.330:
            grade = 'A-'
        elif target > 3:
            grade = 'B+'
        elif target > 2.670:
            grade = 'B'
        elif target > 2.330:
            grade = 'B-'
        elif target > 2:
            grade = 'C+'
            
        return "\tThe student needs a grade >= " + grade + " in: "
    
    return "\tThe student needs a GPA >= %.3f in: " % target

def calculateRequiredGPA(required_GPA, GPA, completed_courses, remaining_courses):
    target = (required_GPA * (getTotalCredits(completed_courses) + getTotalCredits(remaining_courses)) - (GPA * getTotalCredits(completed_courses))) / getTotalCredits(remaining_courses)
    return target

def letterToGPA(letter):
    letter = "".join(findall("[a-zA-Z][\+\-]?", letter))
    letter = letter.upper() # incase DP enters lowercase letter

    if letter == 'A' or letter == 'A+':
        return 4.0
    if letter == 'A-':
        return 3.670
    if letter == 'B+':
        return 3.330
    if letter == 'B':
        return 3.000
    if letter == 'B-':
        return 2.670
    if letter == 'C+':
        return 2.330
    if letter == 'C':
        return 2.000
    if letter == 'F':
        return 0.0
    
    return 'Ignore'

def getGPA(completed_courses):

    gpa = 0.0
    courseCount = 0

    for course in completed_courses:
        grade = letterToGPA(course.grade.split('/')[-1].strip() )
        if (grade == 'Ignore'):
            continue
        
        course_number = "".join(findall('\d+', course.number))
        try:
            credits =int(course_number[1])
        except:
            credits = 3
        
        courseCount = courseCount + credits
        gpa += grade*credits
    
    try:
        return gpa / courseCount
    except Exception as e:
        # DIVIDE BY 0 ERROR?
        return 0

def getTotalCredits(courses):
    sum = 0
    for course in courses:
        if(course.grade):
            grade = letterToGPA(course.grade.split('/')[-1].strip() )
            if (grade == 'Ignore'):
                continue

        course_number = "".join(findall('\d+', course.number))

        try:
            credits =int(course_number[1])
        except:
            credits = 3

        
        sum += credits

    return sum

def is5000Level(course):
    course_number = "".join(findall('\d+', course.number))
    num = int(course_number[0])

    if course.type == 'following':
        if course.semester:
            return True
        else:
            return False
    else:
        return num >= 5

if __name__ == '__main__':
    doAuditMethod('mock')